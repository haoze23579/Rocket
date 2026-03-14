#!/usr/bin/env python
from __future__ import annotations

import argparse
import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree


MATH_TOKEN_RE = re.compile(
    r"(?<!\\)\$\$(.+?)(?<!\\)\$\$|(?<!\\)\$(?!\$)(.+?)(?<!\\)\$(?!\$)",
    re.S,
)
PLACEHOLDER_RE = re.compile(r"(MATH(?:INL|BLK)\d+TOKEN)")


def locate_mml2omml_xsl() -> Path:
    candidates = [
        Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
        Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("MML2OMML.XSL not found in standard Office paths.")


class MdToDocxConverter:
    def __init__(self, input_path: Path, output_path: Path) -> None:
        self.input_path = input_path
        self.output_path = output_path
        self.base_dir = input_path.parent
        self.doc = Document()
        self._setup_styles()
        self._title_done = False

        xsl_path = locate_mml2omml_xsl()
        self._omml_transform = etree.XSLT(etree.parse(str(xsl_path)))
        self._math_placeholders: dict[str, tuple[str, bool]] = {}

    def _protect_math(self, md_text: str) -> str:
        self._math_placeholders.clear()

        def repl(match: re.Match) -> str:
            is_block = match.group(1) is not None
            expr = (match.group(1) if is_block else match.group(2)).strip()
            idx = len(self._math_placeholders)
            token = f"MATHBLK{idx}TOKEN" if is_block else f"MATHINL{idx}TOKEN"
            self._math_placeholders[token] = (expr, is_block)
            if is_block:
                return f"\n\n{token}\n\n"
            return token

        return MATH_TOKEN_RE.sub(repl, md_text)

    def _setup_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        title = self.doc.styles["Title"]
        title.font.name = "Times New Roman"
        title.font.size = Pt(18)
        title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        heading1 = self.doc.styles["Heading 1"]
        heading1.font.name = "Times New Roman"
        heading1.font.size = Pt(16)
        heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        heading2 = self.doc.styles["Heading 2"]
        heading2.font.name = "Times New Roman"
        heading2.font.size = Pt(14)
        heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        heading3 = self.doc.styles["Heading 3"]
        heading3.font.name = "Times New Roman"
        heading3.font.size = Pt(13)
        heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    @staticmethod
    def _style_body_paragraph(paragraph) -> None:
        fmt = paragraph.paragraph_format
        fmt.first_line_indent = Cm(0.74)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(6)

    @staticmethod
    def _style_run(run, bold: bool = False, italic: bool = False, code: bool = False) -> None:
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")

    def _latex_to_omml(self, latex_expr: str):
        expr = latex_expr.strip()
        if not expr:
            return None
        try:
            mathml = latex_to_mathml(expr)
            mathml_root = etree.fromstring(mathml.encode("utf-8"))
            omml = self._omml_transform(mathml_root)
            return parse_xml(etree.tostring(omml, encoding="unicode"))
        except Exception:
            return None

    def _add_text_with_math(
        self,
        paragraph,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        code: bool = False,
    ) -> None:
        if not text:
            return
        if code:
            run = paragraph.add_run(text)
            self._style_run(run, bold=bold, italic=italic, code=True)
            return

        parts = PLACEHOLDER_RE.split(text)
        for part in parts:
            if not part:
                continue
            if part in self._math_placeholders:
                expr, _ = self._math_placeholders[part]
                omml = self._latex_to_omml(expr)
                if omml is not None:
                    paragraph._p.append(omml)
                else:
                    run = paragraph.add_run(expr)
                    self._style_run(run, bold=bold, italic=italic, code=True)
            else:
                run = paragraph.add_run(part.replace(r"\$", "$"))
                self._style_run(run, bold=bold, italic=italic)

    def _append_inline(self, paragraph, node, *, bold=False, italic=False, code=False) -> None:
        if isinstance(node, NavigableString):
            self._add_text_with_math(
                paragraph, str(node), bold=bold, italic=italic, code=code
            )
            return
        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name in ("strong", "b"):
            for child in node.children:
                self._append_inline(
                    paragraph, child, bold=True or bold, italic=italic, code=code
                )
            return
        if name in ("em", "i"):
            for child in node.children:
                self._append_inline(
                    paragraph, child, bold=bold, italic=True or italic, code=code
                )
            return
        if name == "code":
            for child in node.children:
                self._append_inline(
                    paragraph, child, bold=bold, italic=italic, code=True
                )
            return
        if name == "br":
            paragraph.add_run("\n")
            return
        if name == "a":
            text = node.get_text(" ", strip=True)
            href = (node.get("href") or "").strip()
            if href:
                content = f"{text} ({href})" if text and text != href else href
            else:
                content = text
            run = paragraph.add_run(content)
            run.underline = True
            self._style_run(run, bold=bold, italic=italic)
            return
        if name == "img":
            src = (node.get("src") or "").strip()
            alt = (node.get("alt") or "").strip()
            fallback = f"[Image: {alt}] {src}" if alt else f"[Image] {src}"
            run = paragraph.add_run(fallback)
            self._style_run(run, bold=bold, italic=italic)
            return

        for child in node.children:
            self._append_inline(paragraph, child, bold=bold, italic=italic, code=code)

    def _render_image(self, img: Tag) -> None:
        src = (img.get("src") or "").strip()
        alt = (img.get("alt") or "").strip()
        if not src:
            return

        src = src.split("?", 1)[0]
        candidates = [self.base_dir / src, Path.cwd() / src]
        image_path = None
        for c in candidates:
            if c.exists():
                image_path = c
                break

        if image_path:
            self.doc.add_picture(str(image_path), width=Inches(6.0))
            pic_p = self.doc.paragraphs[-1]
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.paragraph_format.space_after = Pt(6)
        else:
            p = self.doc.add_paragraph(f"[Missing image] {src}")
            self._style_body_paragraph(p)

        if alt:
            cap = self.doc.add_paragraph(alt)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(6)

    def _render_list_as_paragraph(self, list_tag: Tag) -> None:
        items = list_tag.find_all("li", recursive=False)
        if not items:
            return

        p = self.doc.add_paragraph()
        self._style_body_paragraph(p)

        for idx, li in enumerate(items, start=1):
            if idx > 1:
                self._add_text_with_math(p, "；")
            self._add_text_with_math(p, f"（{idx}）")

            for child in li.contents:
                if isinstance(child, Tag) and child.name.lower() in ("ul", "ol"):
                    continue
                if isinstance(child, Tag) and child.name.lower() == "p":
                    for gc in child.children:
                        self._append_inline(p, gc)
                else:
                    self._append_inline(p, child)

        self._add_text_with_math(p, "。")

    def _render_table(self, table_tag: Tag) -> None:
        rows = table_tag.find_all("tr")
        if not rows:
            return
        col_count = max(len(r.find_all(["th", "td"], recursive=False)) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"], recursive=False)
            for j, cell in enumerate(cells):
                p = table.cell(i, j).paragraphs[0]
                p.text = ""
                for child in cell.children:
                    self._append_inline(p, child)
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)

    def _render_heading(self, tag: Tag) -> None:
        level = int(tag.name[1])
        text = tag.get_text(" ", strip=True)
        if not text:
            return

        if not self._title_done and level == 1:
            p = self.doc.add_paragraph(text, style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            self._title_done = True
            return

        mapped_level = max(1, level - 1 if self._title_done else level)
        mapped_level = min(mapped_level, 3)
        p = self.doc.add_heading("", level=mapped_level)
        for child in tag.children:
            self._append_inline(p, child)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)

    def _render_block_equation(self, latex_expr: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        omml = self._latex_to_omml(latex_expr)
        if omml is not None:
            p._p.append(omml)
        else:
            run = p.add_run(latex_expr)
            self._style_run(run, code=True)

    def _render_paragraph(self, tag: Tag) -> None:
        children = [
            c
            for c in tag.contents
            if not (isinstance(c, NavigableString) and not str(c).strip())
        ]
        if (
            len(children) == 1
            and isinstance(children[0], Tag)
            and children[0].name.lower() == "img"
        ):
            self._render_image(children[0])
            return

        text_content = tag.get_text("", strip=False)
        token = text_content.strip()
        if (
            token in self._math_placeholders
            and self._math_placeholders[token][1]
            and len(token.split()) == 1
        ):
            self._render_block_equation(self._math_placeholders[token][0])
            return

        p = self.doc.add_paragraph()
        self._style_body_paragraph(p)
        for child in tag.children:
            self._append_inline(p, child)

    def _render_pre(self, tag: Tag) -> None:
        code_text = tag.get_text("\n")
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(code_text)
        self._style_run(run, code=True)

    def _render_block(self, node) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                p = self.doc.add_paragraph()
                self._style_body_paragraph(p)
                self._add_text_with_math(p, text)
            return
        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._render_heading(node)
            return
        if name == "p":
            self._render_paragraph(node)
            return
        if name in ("ul", "ol"):
            self._render_list_as_paragraph(node)
            return
        if name == "table":
            self._render_table(node)
            return
        if name == "pre":
            self._render_pre(node)
            return
        if name == "blockquote":
            p = self.doc.add_paragraph(node.get_text(" ", strip=True))
            self._style_body_paragraph(p)
            p.paragraph_format.left_indent = Cm(0.74)
            return
        if name == "hr":
            self.doc.add_paragraph("")
            return
        if name == "img":
            self._render_image(node)
            return

        for child in node.children:
            self._render_block(child)

    def convert(self) -> None:
        md_text = self.input_path.read_text(encoding="utf-8")
        protected_md = self._protect_math(md_text)
        html = markdown.markdown(
            protected_md,
            extensions=[
                "extra",
                "tables",
                "fenced_code",
                "sane_lists",
            ],
        )
        soup = BeautifulSoup(html, "html.parser")
        for top in soup.contents:
            self._render_block(top)

        self.doc.save(self.output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert paper markdown to thesis-like DOCX with equation rendering."
    )
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("paper.md"),
        help="Input markdown file.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("paper.docx"),
        help="Output docx file.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    converter = MdToDocxConverter(args.input, args.output)
    converter.convert()
    print(f"WROTE: {args.output.resolve()}")


if __name__ == "__main__":
    main()
